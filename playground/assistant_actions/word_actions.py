import time
from docx import Document
from docx.shared import Inches

from playground.actions_manager import agent_action

import zipfile
import os

from playground.global_values import GlobalValues


@agent_action
def extract_text_from_doc(doc_filename):
    """
    Extract all text from a Word document.

    Parameters:
    doc_filename (str): The path to the Word document (.docx).

    Returns:
    str: The extracted text from the document.
    """
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc = Document(doc_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


@agent_action
def extract_all_from_doc(doc_filename, output_folder):
    """
    Extract all text, images, and code from a Word document.

    Parameters:
    doc_filename (str): Path to the Word document (.docx).
    output_folder (str): Directory to save the extracted images and code files.

    Returns:
    str: Full document text with placeholders for images and code.
    list: List of paths to the extracted images.
    list: List of paths to the extracted code files.
    """
    # Ensure the output directory exists
    output_folder = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, output_folder)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Open the document
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc = Document(doc_path)

    full_text = []
    image_paths = []
    code_paths = []
    image_counter = 1
    code_counter = 1

    # Iterate through the document elements
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        # Check if the paragraph contains code
        if "```" in paragraph_text:
            code_block = paragraph_text.split("```")[1]
            code_filename = os.path.join(output_folder, f"code_block_{code_counter}.py")
            with open(code_filename, "w") as code_file:
                code_file.write(code_block)
            code_paths.append(code_filename)
            full_text.append(f"[CODE BLOCK {code_counter}]")
            code_counter += 1
        else:
            full_text.append(paragraph_text)

        for run in paragraph.runs:
            if "drawing" in run._element.xml:
                drawing_element = run._element.xpath(".//a:blip/@r:embed")
                if drawing_element:
                    image_id = drawing_element[0]
                    # Placeholder for image
                    full_text.append(f" [image{image_counter}.png] ")
                    image_counter += 1

    # Extract images from the docx file
    with zipfile.ZipFile(doc_path, "r") as docx_zip:
        for file in docx_zip.namelist():
            if file.startswith("word/media/"):
                image_data = docx_zip.read(file)
                image_filename = os.path.basename(file)
                image_path = os.path.join(output_folder, image_filename)
                with open(image_path, "wb") as image_file:
                    image_file.write(image_data)
                image_paths.append(image_path)

    full_document_text = "\n".join(full_text)
    return full_document_text, image_paths, code_paths


# @agent_action
# def create_word_doc(doc_filename):
#     """
#     Create a new Word document.

#     Parameters:
#     doc_filename (str): The path to the Word document.
#     """
#     doc = Document()
#     doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
#     doc.save(doc_path)
#     return "Document created successfully."


@agent_action
def create_word_doc(doc_filename):
    """
    Create a new Word document.

    Parameters:
    doc_filename (str): The path to the Word document.
    """
    doc = Document()
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc.save(doc_path)

    # Update file creation and modification times
    current_time = time.time()
    os.utime(doc_path, (current_time, current_time))
    return "Document created successfully."


@agent_action
def append_paragraph_to_word_doc(doc_filename, paragraph_text):
    """
    Append a paragraph to a Word document.

    Parameters:
    doc_filename (str): The path to the Word document.
    paragraph_text (str): The text to append as a paragraph.
    """
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc = Document(doc_path)
    doc.add_paragraph(paragraph_text)
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc.save(doc_path)
    return "Paragraph added successfully."


@agent_action
def append_heading_to_word_doc(doc_filename, heading_text, level):
    """
    Append a heading to a Word document.

    Parameters:
    doc_filename (str): The path to the Word document.
    heading_text (str): The text to append as a heading.
    level (int): The heading level (0-9).
    """
    level = int(level)
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc = Document(doc_path)
    doc.add_heading(heading_text, level=level)
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc.save(doc_path)
    return "Heading added successfully."


def normalize_path(path):
    """
    Normalize the given path, removing any reference indicators like '.' or '..'.

    Parameters:
    path (str): The path to normalize.

    Returns:
    str: The normalized path.
    """
    return os.path.normpath(path)


@agent_action
def append_image_to_word_doc(doc_filename, image_filename, caption_text):
    """
    Append an image with a caption to a Word document.

    Parameters:
    doc_filename (str): The path to the Word document.
    image_filename (str): The path to the image file to append.
    caption_text (str): The caption text to append below the image.
    """
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc = Document(doc_path)
    image_path = os.path.join(
        GlobalValues.ASSISTANTS_WORKING_FOLDER, normalize_path(image_filename)
    )
    doc.add_picture(image_path, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = 1  # Center the image
    caption = doc.add_paragraph(caption_text)
    caption.alignment = 1  # Center the caption
    caption.style = "Caption"
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc.save(doc_path)
    return "Image added successfully."


@agent_action
def append_code_to_word_doc(doc_filename, code_text):
    """
    Append a code block to a Word document.

    Parameters:
    doc_filename (str): The path to the Word document.
    code_text (str): The code text to append.
    """
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc = Document(doc_path)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Courier New"
    doc_path = os.path.join(GlobalValues.ASSISTANTS_WORKING_FOLDER, doc_filename)
    doc.save(doc_path)
    return "Code block added successfully."


# # Usage example:
# # Create a new document first
# doc = Document()
# doc.save('example.docx')

# # Append content to the document
# append_paragraph_to_word_doc('example.docx', 'This is a sample paragraph.')
# append_heading_to_word_doc('example.docx', 'Sample Heading', 2)
# append_image_to_word_doc('example.docx', '/path/to/image.png', 'Sample Image Caption')
# append_code_to_word_doc('example.docx', 'print("Hello, world!")')

# print("Content added to the document successfully.")
